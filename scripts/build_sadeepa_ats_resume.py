from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, black
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(r"D:\ic-marketplace")
DOCX_OUTPUT = ROOT / "output" / "docx" / "Sadeepa_Amaranayake_ATS_Resume.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "Sadeepa_Amaranayake_ATS_Resume.pdf"


PROFILE = {
    "name": "SADEEPA AMARANAYAKE",
    "target": "Software Engineering Intern | Full-Stack Developer Intern",
    "phone": "+94 75 441 4937",
    "email": "sadeepaamaranayake@gmail.com",
    "github": "github.com/sadeepaghost",
    "linkedin": "linkedin.com/in/sadeepa-amaranayake-b123b0292",
    "portfolio": "sadeepa-portfolio.netlify.app",
    "summary": (
        "Full-stack developer and Bachelor of Information Technology undergraduate specializing in "
        "Software Systems. Builds web applications with JavaScript, TypeScript, React, Node.js, "
        "Express.js, Laravel, PHP, MongoDB, and MySQL. Hands-on experience with REST APIs, JWT "
        "authentication, role-based access control (RBAC), MongoDB aggregation, Redis caching, "
        "Stripe and OpenAI API integrations, Docker, and GitHub Actions."
    ),
    "skills": [
        ("Programming Languages", "JavaScript, TypeScript, PHP, Python, Java, Dart, C"),
        ("Frontend", "React.js, Tailwind CSS, Blade, Livewire, Alpine.js, Flutter"),
        ("Backend and APIs", "Node.js, Express.js, Laravel, REST APIs, JWT Authentication, RBAC"),
        ("Databases", "MongoDB, MySQL, Redis"),
        ("Cloud and DevOps", "AWS, Docker, GitHub Actions, Git, GitHub"),
        ("Tools and Concepts", "Postman, VS Code, OOP, Data Structures, MVC, CRUD, API Integration"),
    ],
    "projects": [
        {
            "name": "Multi-Vendor E-Commerce Marketplace",
            "year": "2026",
            "tech": "MongoDB, Express.js, React.js, Node.js, Redis, Stripe",
            "bullets": [
                "Built a multi-vendor e-commerce platform with independent vendor storefronts and RBAC for customer, vendor, and administrator roles.",
                "Implemented MongoDB data models and Redis caching to support inventory synchronization across multiple vendors.",
                "Created REST API integrations and connected Stripe payment processing to a React single-page application.",
            ],
            "url": "github.com/sadeepaghost/eCommerce-marketplace",
        },
        {
            "name": "IC Parts Marketplace",
            "year": "2026",
            "tech": "Laravel, Blade, Tailwind CSS, Livewire, Alpine.js, OpenAI API",
            "bullets": [
                "Developed an electronics-components marketplace with catalog search and automated quote-request workflows.",
                "Integrated the OpenAI API to assist part matching, sourcing, and quote preparation.",
                "Built real-time filtered search and inventory updates with Livewire and Alpine.js.",
            ],
            "url": "github.com/sadeepaghost/IC-Marketplace",
        },
        {
            "name": "University Admin System - Dashboard Module",
            "year": "2026",
            "tech": "MongoDB, Express.js, React.js, Node.js",
            "bullets": [
                "Developed the dashboard module for a team-based university administration system under shared schema and code-ownership rules.",
                "Created Express.js endpoints and MongoDB aggregation pipelines, including lookup joins, to report live student, course, and lecturer counts.",
                "Diagnosed routing, data-layer, and shared frontend integration issues while coordinating changes with teammates.",
            ],
            "url": None,
        },
    ],
    "education": {
        "degree": "Bachelor of Information Technology - Software Systems",
        "institution": "University of Kelaniya",
        "dates": "2023 - Present | Expected Graduation: 2027",
    },
    "certifications": [
        "The Complete Python Bootcamp: Zero to Expert - Udemy, 2026",
        "JavaScript: Beginner to Web Developer - Udemy, 2026",
        "Blockchain and Web3 Mastery - Udemy, 2025",
    ],
}


# compact_reference_guide with a named ATS Resume Override:
# Letter page; 0.65-inch margins; Arial; black text; single-column body;
# compact, explicit spacing; no tables, text boxes, icons, headers, or footers.
ATS_MARGIN_IN = 0.65
FONT_NAME = "Arial"
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x35, 0x35, 0x35)


def set_run_font(run, *, size: float, bold: bool = False, color: RGBColor = INK):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_free_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(ATS_MARGIN_IN)
    section.bottom_margin = Inches(ATS_MARGIN_IN)
    section.left_margin = Inches(ATS_MARGIN_IN)
    section.right_margin = Inches(ATS_MARGIN_IN)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)


def configure_docx_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(9.0)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.0)
    normal.paragraph_format.line_spacing = 1.0

    heading = styles["Heading 1"]
    heading.font.name = FONT_NAME
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    heading.font.size = Pt(10.4)
    heading.font.bold = True
    heading.font.color.rgb = INK
    heading.paragraph_format.space_before = Pt(5.4)
    heading.paragraph_format.space_after = Pt(2.1)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True

    custom_styles = {
        "Resume Name": (20.5, True, 0, 0.2, 1.0),
        "Resume Target": (10.7, True, 0, 2.2, 1.0),
        "Resume Contact": (8.2, False, 0, 0.6, 1.0),
        "Resume Summary": (9.25, False, 0, 1.2, 1.0),
        "Resume Skill": (8.9, False, 0, 0.25, 1.0),
        "Resume Project": (9.6, True, 3.4, 0.15, 1.0),
        "Resume Tech": (8.55, False, 0, 0.15, 1.0),
        "Resume Bullet": (9.05, False, 0, 0.35, 1.0),
        "Resume Link": (8.25, False, 0, 0.5, 1.0),
        "Resume Education": (9.1, False, 0, 0.35, 1.0),
    }

    for name, (size, bold, before, after, line) in custom_styles.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = MUTED if name in {"Resume Contact", "Resume Tech", "Resume Link"} else INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line
        if name in {"Resume Name", "Resume Target", "Resume Contact"}:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name == "Resume Project":
            style.paragraph_format.keep_with_next = True


def add_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "252")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "252")
    ind.set(qn("w:hanging"), "126")
    p_pr.append(ind)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_pr.append(fonts)
    lvl.append(r_pr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 8.2, bold: bool = False):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_el = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "353535")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz_cs)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run_el.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run_el.append(text_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def add_label_value_paragraph(doc: Document, label: str, value: str, style: str):
    p = doc.add_paragraph(style=style)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=8.9, bold=True)
    value_run = p.add_run(value)
    set_run_font(value_run, size=8.9)
    return p


def build_docx():
    DOCX_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    set_cell_free_page(section)
    configure_docx_styles(doc)
    bullet_num_id = add_bullet_numbering(doc)

    doc.core_properties.title = "Sadeepa Amaranayake - ATS Resume"
    doc.core_properties.subject = "Software Engineering and Full-Stack Development Internship Resume"
    doc.core_properties.author = "Sadeepa Amaranayake"
    doc.core_properties.keywords = (
        "Software Engineering Intern, Full-Stack Developer, React, Node.js, Express.js, Laravel, "
        "JavaScript, TypeScript, MongoDB, MySQL, Redis, REST API, Docker, AWS"
    )

    p = doc.add_paragraph(style="Resume Name")
    p.add_run(PROFILE["name"])

    p = doc.add_paragraph(style="Resume Target")
    p.add_run(PROFILE["target"])

    p = doc.add_paragraph(style="Resume Contact")
    p.add_run("Phone: ")
    add_hyperlink(p, PROFILE["phone"], f"tel:{PROFILE['phone'].replace(' ', '')}")
    p.add_run(" | Email: ")
    add_hyperlink(p, PROFILE["email"], f"mailto:{PROFILE['email']}")

    p = doc.add_paragraph(style="Resume Contact")
    p.add_run("GitHub: ")
    add_hyperlink(p, PROFILE["github"], f"https://{PROFILE['github']}")
    p.add_run(" | LinkedIn: ")
    add_hyperlink(p, PROFILE["linkedin"], f"https://{PROFILE['linkedin']}")
    p.add_run(" | Portfolio: ")
    add_hyperlink(p, PROFILE["portfolio"], f"https://{PROFILE['portfolio']}")

    doc.add_paragraph("PROFESSIONAL SUMMARY", style="Heading 1")
    doc.add_paragraph(PROFILE["summary"], style="Resume Summary")

    doc.add_paragraph("TECHNICAL SKILLS", style="Heading 1")
    for label, value in PROFILE["skills"]:
        add_label_value_paragraph(doc, label, value, "Resume Skill")

    doc.add_paragraph("PROJECT EXPERIENCE", style="Heading 1")
    for project in PROFILE["projects"]:
        p = doc.add_paragraph(style="Resume Project")
        title = p.add_run(f"{project['name']} | {project['year']}")
        set_run_font(title, size=9.6, bold=True)

        p = doc.add_paragraph(style="Resume Tech")
        label = p.add_run("Technologies: ")
        set_run_font(label, size=8.55, bold=True, color=MUTED)
        tech = p.add_run(project["tech"])
        set_run_font(tech, size=8.55, color=MUTED)

        for bullet in project["bullets"]:
            p = doc.add_paragraph(bullet, style="Resume Bullet")
            apply_num(p, bullet_num_id)

        if project["url"]:
            p = doc.add_paragraph(style="Resume Link")
            label = p.add_run("GitHub: ")
            set_run_font(label, size=8.25, bold=True, color=MUTED)
            add_hyperlink(p, project["url"], f"https://{project['url']}", size=8.25)

    doc.add_paragraph("EDUCATION", style="Heading 1")
    p = doc.add_paragraph(style="Resume Education")
    degree = p.add_run(PROFILE["education"]["degree"])
    set_run_font(degree, size=9.1, bold=True)
    p.add_run(f" | {PROFILE['education']['institution']}")
    p = doc.add_paragraph(PROFILE["education"]["dates"], style="Resume Education")

    doc.add_paragraph("CERTIFICATIONS", style="Heading 1")
    for certification in PROFILE["certifications"]:
        p = doc.add_paragraph(certification, style="Resume Bullet")
        apply_num(p, bullet_num_id)

    doc.save(DOCX_OUTPUT)


def pdf_link(text: str, url: str) -> str:
    return f'<a href="{escape(url)}" color="#353535"><u>{escape(text)}</u></a>'


def make_pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "ResumeName",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=20.5,
            leading=21.8,
            textColor=black,
            alignment=TA_CENTER,
            spaceAfter=0,
        ),
        "target": ParagraphStyle(
            "ResumeTarget",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.7,
            leading=11.8,
            textColor=black,
            alignment=TA_CENTER,
            spaceAfter=1.5,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.2,
            leading=9.5,
            textColor=HexColor("#353535"),
            alignment=TA_CENTER,
            spaceAfter=0.25,
        ),
        "heading": ParagraphStyle(
            "ResumeHeading",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.4,
            leading=11.6,
            textColor=black,
            alignment=TA_LEFT,
            spaceBefore=5.2,
            spaceAfter=2.0,
            keepWithNext=True,
        ),
        "summary": ParagraphStyle(
            "ResumeSummary",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.25,
            leading=11.1,
            textColor=black,
            alignment=TA_LEFT,
            spaceAfter=0.5,
        ),
        "skill": ParagraphStyle(
            "ResumeSkill",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.9,
            leading=10.4,
            textColor=black,
            alignment=TA_LEFT,
            spaceAfter=0.05,
        ),
        "project": ParagraphStyle(
            "ResumeProject",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.6,
            leading=11.1,
            textColor=black,
            alignment=TA_LEFT,
            spaceBefore=3.4,
            spaceAfter=0.1,
            keepWithNext=True,
        ),
        "tech": ParagraphStyle(
            "ResumeTech",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.55,
            leading=10.0,
            textColor=HexColor("#353535"),
            alignment=TA_LEFT,
            spaceAfter=0,
            keepWithNext=True,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.05,
            leading=10.6,
            textColor=black,
            alignment=TA_LEFT,
            leftIndent=10,
            bulletIndent=1,
            firstLineIndent=0,
            spaceAfter=0.15,
            bulletFontName="Helvetica",
            bulletFontSize=9.05,
        ),
        "link": ParagraphStyle(
            "ResumeLink",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.25,
            leading=9.5,
            textColor=HexColor("#353535"),
            alignment=TA_LEFT,
            spaceAfter=0.3,
        ),
        "education": ParagraphStyle(
            "ResumeEducation",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=10.5,
            textColor=black,
            alignment=TA_LEFT,
            spaceAfter=0.2,
        ),
    }


def build_pdf():
    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    styles = make_pdf_styles()
    pdf = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        leftMargin=ATS_MARGIN_IN * inch,
        rightMargin=ATS_MARGIN_IN * inch,
        topMargin=ATS_MARGIN_IN * inch,
        bottomMargin=ATS_MARGIN_IN * inch,
        title="Sadeepa Amaranayake - ATS Resume",
        author="Sadeepa Amaranayake",
        subject="Software Engineering and Full-Stack Development Internship Resume",
    )

    story = [
        Paragraph(escape(PROFILE["name"]), styles["name"]),
        Paragraph(escape(PROFILE["target"]), styles["target"]),
        Paragraph(
            f'Phone: {pdf_link(PROFILE["phone"], "tel:" + PROFILE["phone"].replace(" ", ""))}'
            f' | Email: {pdf_link(PROFILE["email"], "mailto:" + PROFILE["email"])}',
            styles["contact"],
        ),
        Paragraph(
            f'GitHub: {pdf_link(PROFILE["github"], "https://" + PROFILE["github"])}'
            f' | LinkedIn: {pdf_link(PROFILE["linkedin"], "https://" + PROFILE["linkedin"])}'
            f' | Portfolio: {pdf_link(PROFILE["portfolio"], "https://" + PROFILE["portfolio"])}',
            styles["contact"],
        ),
        Paragraph("PROFESSIONAL SUMMARY", styles["heading"]),
        Paragraph(escape(PROFILE["summary"]), styles["summary"]),
        Paragraph("TECHNICAL SKILLS", styles["heading"]),
    ]

    for label, value in PROFILE["skills"]:
        story.append(Paragraph(f"<b>{escape(label)}:</b> {escape(value)}", styles["skill"]))

    story.append(Paragraph("PROJECT EXPERIENCE", styles["heading"]))
    for project in PROFILE["projects"]:
        block = [
            Paragraph(f"{escape(project['name'])} | {escape(project['year'])}", styles["project"]),
            Paragraph(f"<b>Technologies:</b> {escape(project['tech'])}", styles["tech"]),
        ]
        for bullet in project["bullets"]:
            block.append(Paragraph(escape(bullet), styles["bullet"], bulletText="-"))
        if project["url"]:
            block.append(
                Paragraph(
                    f'<b>GitHub:</b> {pdf_link(project["url"], "https://" + project["url"])}',
                    styles["link"],
                )
            )
        story.append(KeepTogether(block))

    story.extend(
        [
            Paragraph("EDUCATION", styles["heading"]),
            Paragraph(
                f'<b>{escape(PROFILE["education"]["degree"])}</b> | '
                f'{escape(PROFILE["education"]["institution"])}',
                styles["education"],
            ),
            Paragraph(escape(PROFILE["education"]["dates"]), styles["education"]),
            Paragraph("CERTIFICATIONS", styles["heading"]),
        ]
    )
    for certification in PROFILE["certifications"]:
        story.append(Paragraph(escape(certification), styles["bullet"], bulletText="-"))

    pdf.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUTPUT)
    print(PDF_OUTPUT)
